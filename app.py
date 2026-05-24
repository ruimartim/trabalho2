import streamlit as st
import io
import re

st.set_page_config(page_title="Pipeline de Texto", page_icon="🌿")

st.markdown("""
<style>
    .stApp { background-color: #f5f5f0 !important; }
    h1, h2, h3 { color: #4a6741 !important; }
    .stButton>button { background-color: #c8b89a !important; color: white !important; border: none; }
    .stTextArea textarea { border-color: #8faa8b !important; }
</style>
""", unsafe_allow_html=True)

st.title("Pipeline de Pré-Processamento de Texto")
st.write("Trabalho Prático 2 - Laboratório de Programação")
st.divider()

try:
    import fitz
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

def extrair_texto(ficheiro):
    nome = ficheiro.name.lower()
    dados = ficheiro.read()
    if nome.endswith(".txt"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return dados.decode(enc)
            except:
                continue
        return dados.decode("utf-8", errors="replace")
    elif nome.endswith(".pdf"):
        if not PDF_OK:
            return "ERRO: instala pymupdf com: pip install pymupdf"
        doc = fitz.open(stream=dados, filetype="pdf")
        return "\n".join(p.get_text() for p in doc)
    elif nome.endswith(".docx"):
        if not DOCX_OK:
            return "ERRO: instala python-docx com: pip install python-docx"
        doc = Document(io.BytesIO(dados))
        return "\n".join(p.text for p in doc.paragraphs)
    return "ERRO: formato não suportado."

def limpar_texto(texto, cfg):
    t = texto
    if cfg["artefactos"]:
        t = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\x80-\xFF]', '', t)
        t = re.sub(r'[_\-=~]{4,}', '', t)
    if cfg["encoding"]:
        subs = {'\u2019': "'", '\u201c': '"', '\u201d': '"', '\u2013': '-', '\u2014': '--'}
        for k, v in subs.items():
            t = t.replace(k, v)
    if cfg["cabecalhos"]:
        from collections import Counter
        linhas = t.split('\n')
        contagem = Counter(l.strip() for l in linhas if l.strip())
        repetidas = {l for l, c in contagem.items() if c >= 3 and len(l) < 120}
        t = '\n'.join(l for l in linhas if l.strip() not in repetidas)
    if cfg["quebras"]:
        t = re.sub(r'(?<![.!?:])\n(?!\n)', ' ', t)
    if cfg["paragrafos"]:
        t = re.sub(r'\n{3,}', '\n\n', t)
    if cfg["espacos"]:
        t = '\n'.join(re.sub(r'[ \t]+', ' ', l).strip() for l in t.split('\n'))
    return t.strip()

def detectar_idioma(texto):
    amostra = texto[:2000].lower()
    palavras = re.findall(r'\b\w+\b', amostra)
    pt = sum(1 for w in palavras if w in ['de','que','em','para','com','uma','por','não','como','mais'])
    en = sum(1 for w in palavras if w in ['the','of','and','to','in','is','it','for','on','with'])
    return "Português" if pt >= en else "English"

def dividir_chunks(texto, tamanho):
    paragrafos = re.split(r'\n{2,}', texto)
    chunks, atual = [], ""
    for p in paragrafos:
        if len(atual) + len(p) + 2 <= tamanho:
            atual += ("\n\n" if atual else "") + p
        else:
            if atual:
                chunks.append(atual.strip())
            atual = p
    if atual:
        chunks.append(atual.strip())
    return [c for c in chunks if c]

def gerar_prompt(chunk, idioma, tipo):
    tipos = {
        "Correção gramatical": f"The text is in {idioma}. Correct grammar and punctuation. Return only the corrected text.\n\nText:\n{chunk}",
        "Resumo": f"The text is in {idioma}. Write a short summary. Return only the summary.\n\nText:\n{chunk}",
        "Estruturação": f"The text is in {idioma}. Organize into clear paragraphs. Return only the restructured text.\n\nText:\n{chunk}",
    }
    return tipos[tipo]

# Sidebar
st.sidebar.header("Opções da Pipeline")
cfg = {
    "artefactos": st.sidebar.checkbox("Remover artefactos", value=True),
    "encoding":   st.sidebar.checkbox("Corrigir encoding", value=True),
    "cabecalhos": st.sidebar.checkbox("Remover cabeçalhos repetidos", value=True),
    "quebras":    st.sidebar.checkbox("Corrigir quebras de linha", value=True),
    "paragrafos": st.sidebar.checkbox("Reconstruir parágrafos", value=True),
    "espacos":    st.sidebar.checkbox("Normalizar espaços", value=True),
}
st.sidebar.divider()
tamanho_chunk = st.sidebar.slider("Tamanho do chunk", 300, 3000, 1000, 100)
tipo_prompt = st.sidebar.selectbox("Tipo de prompt", ["Correção gramatical", "Resumo", "Estruturação"])

# Etapa 1
st.header("Etapa 1 - Upload de Ficheiro")
ficheiro = st.file_uploader("Escolhe um ficheiro PDF, DOCX ou TXT", type=["pdf", "docx", "txt"])

if ficheiro:
    st.session_state["bruto"] = extrair_texto(ficheiro)
    st.success("Ficheiro carregado com sucesso!")

if "bruto" in st.session_state:
    st.text_area("Texto extraído (bruto)", st.session_state["bruto"], height=200)
    st.write(f"Caracteres: {len(st.session_state['bruto'])} | Palavras: {len(st.session_state['bruto'].split())}")
    st.divider()

    # Etapa 2
    st.header("Etapa 2 - Limpeza do Texto")

    if st.button("Executar Limpeza"):
        st.session_state["limpo"] = limpar_texto(st.session_state["bruto"], cfg)

    if "limpo" in st.session_state:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Antes")
            st.text_area("", st.session_state["bruto"], height=200, key="antes")
            st.write(f"{len(st.session_state['bruto'])} caracteres")
        with col2:
            st.subheader("Depois")
            st.text_area("", st.session_state["limpo"], height=200, key="depois")
            st.write(f"{len(st.session_state['limpo'])} caracteres")
        st.divider()


        # Etapa 3
        st.header("Etapa 3 - Preparação para o Modelo")
        idioma = detectar_idioma(st.session_state["limpo"])
        st.write(f"Idioma detetado: **{idioma}**")
        chunks = dividir_chunks(st.session_state["limpo"], tamanho_chunk)
        st.write(f"Número de chunks: **{len(chunks)}**")

        for i, chunk in enumerate(chunks):
            with st.expander(f"Chunk {i+1} ({len(chunk)} caracteres)"):
                t1, t2 = st.tabs(["Texto", "Prompt"])
                with t1:
                    st.text_area("", chunk, height=150, key=f"c{i}")
                with t2:
                    st.text_area("", gerar_prompt(chunk, idioma, tipo_prompt), height=180, key=f"p{i}")

        st.divider()

        # Etapa 4
        st.header("Etapa 4 - Enviar para o Modelo")
        if st.button("Enviar para o SLM"):
            import requests
            respostas = []
            cancelado = False
            progress = st.progress(0)
            status = st.empty()
            parar = st.button("Cancelar")

            for i, chunk in enumerate(chunks):
                if parar:
                    cancelado = True
                    st.warning("Operação cancelada.")
                    break
                status.write(f"A processar chunk {i+1} de {len(chunks)}...")
                progress.progress((i+1) / len(chunks))
                prompt = gerar_prompt(chunk, idioma, tipo_prompt)
                try:
                    resposta = requests.post(
                        "https://reality.utad.net/slm",
                        json={
                            "model": "llama-3.2-1b-instruct",
                            "messages": [{"role": "user", "content": prompt}]
                        },
                        timeout=300
                    )
                    texto_resposta = resposta.json()["choices"][0]["message"]["content"]
                    respostas.append(texto_resposta)
                except requests.exceptions.Timeout:
                    st.warning(f"Chunk {i+1} demorou mais de 5 minutos.")
                    opcao = st.radio("O que queres fazer?", ["Continuar", "Terminar"], key=f"opcao_{i}")
                    if opcao == "Terminar":
                        cancelado = True
                        break
                    else:
                        respostas.append("(chunk ignorado por timeout)")
                except Exception as e:
                    respostas.append(f"ERRO: {str(e)}")

            if not cancelado:
                status.write("Concluído!")
            st.session_state["respostas"] = respostas

        if "respostas" in st.session_state:
            for i, r in enumerate(st.session_state["respostas"]):
                with st.expander(f"Resposta chunk {i+1}"):
                    st.write(r)

        st.divider()

        # Etapa 5
        st.header("Etapa 5 - Relatório")
        if st.button("Gerar Relatório"):
            st.session_state["relatorio"] = f"""Parâmetros da pipeline:
- Remover artefactos: {cfg['artefactos']}
- Corrigir encoding: {cfg['encoding']}
- Remover cabeçalhos: {cfg['cabecalhos']}
- Corrigir quebras de linha: {cfg['quebras']}
- Reconstruir parágrafos: {cfg['paragrafos']}
- Normalizar espaços: {cfg['espacos']}
- Tamanho do chunk: {tamanho_chunk}
- Tipo de prompt: {tipo_prompt}
- Idioma: {idioma}

Texto antes:
{st.session_state['bruto']}

Texto depois:
{st.session_state['limpo']}

Avaliação:
- Caracteres antes: {len(st.session_state['bruto'])}
- Caracteres depois: {len(st.session_state['limpo'])}
- Removidos: {len(st.session_state['bruto']) - len(st.session_state['limpo'])}
- Chunks: {len(chunks)}"""

        if "relatorio" in st.session_state:
            st.text_area("Relatório", st.session_state["relatorio"], height=300)
            formato = st.selectbox("Formato do relatório", ["TXT", "PDF", "DOCX"])

            if formato == "TXT":
                st.download_button("Descarregar Relatório", st.session_state["relatorio"].encode(), "relatorio.txt")

            elif formato == "PDF":
                try:
                    from fpdf import FPDF
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Helvetica", size=12)
                    for linha in st.session_state["relatorio"].split('\n'):
                        linha_limpa = linha.encode('latin-1', 'replace').decode('latin-1')
                        pdf.cell(0, 10, linha_limpa, ln=True)
                    st.download_button("Descarregar Relatório", bytes(pdf.output()), "relatorio.pdf")
                except ImportError:
                    st.error("Instala fpdf2 com: pip install fpdf2")

            elif formato == "DOCX":
                try:
                    from docx import Document
                    doc = Document()
                    for linha in st.session_state["relatorio"].split('\n'):
                        doc.add_paragraph(linha)
                    buf = io.BytesIO()
                    doc.save(buf)
                    st.download_button("Descarregar Relatório", buf.getvalue(), "relatorio.docx")
                except ImportError:
                    st.error("Instala python-docx com: pip install python-docx")

        st.divider()
        st.download_button("Descarregar texto limpo", st.session_state["limpo"].encode(), "texto_limpo.txt")